"""读取 Word 文档"""
from docx import Document

path = r"D:\my_crew\knowledge\8255 E01售后黑卡闪专项问题跟进分析.docx"
doc = Document(path)

print("=" * 60)
print("PARAGRAPHS:")
print("=" * 60)
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"[P{i}] {text}")
        print()

print("\n" + "=" * 60)
print(f"TABLES: {len(doc.tables)}")
print("=" * 60)
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip()[:100] for cell in row.cells]
        print(f"  Row {ri}: {cells}")
